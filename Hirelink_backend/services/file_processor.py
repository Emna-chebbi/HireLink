"""
File Processor for extracting text from various resume formats
"""

import os
import re
from typing import Optional, Tuple
import PyPDF2
from docx import Document
from io import BytesIO
import warnings
warnings.filterwarnings('ignore')

class ResumeFileProcessor:
    """
    Process uploaded resume files and extract text
    """
    
    @staticmethod
    def extract_text_from_file(file, filename: str) -> Tuple[str, dict]:
        """
        Extract text from uploaded file
        
        Args:
            file: Django uploaded file object
            filename: Original filename
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {
            'filename': filename,
            'file_type': filename.split('.')[-1].lower() if '.' in filename else 'unknown',
            'success': False,
            'error': None
        }
        
        try:
            # Read file content
            content = file.read()
            
            # Determine file type and extract text
            if filename.lower().endswith('.pdf'):
                text = ResumeFileProcessor._extract_from_pdf(content)
            elif filename.lower().endswith(('.docx', '.doc')):
                text = ResumeFileProcessor._extract_from_docx(content)
            elif filename.lower().endswith('.txt'):
                text = ResumeFileProcessor._extract_from_txt(content)
            else:
                # Try to extract as text for unknown formats
                try:
                    text = content.decode('utf-8', errors='ignore')
                except:
                    text = str(content, errors='ignore')
            
            # Clean the extracted text
            text = ResumeFileProcessor._clean_text(text)
            
            # Update metadata
            metadata.update({
                'success': True,
                'word_count': len(text.split()),
                'char_count': len(text),
                'file_size_kb': len(content) / 1024
            })
            
            return text, metadata
            
        except Exception as e:
            metadata['error'] = str(e)
            return "", metadata
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            # Try PyPDF2 first
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
            # Fallback: Try pdfminer
            try:
                from pdfminer.high_level import extract_text
                pdf_file = BytesIO(content)
                text = extract_text(pdf_file)
            except:
                text = "[PDF content - extraction failed]"
        
        return text
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return "[DOCX content - extraction failed]"
    
    @staticmethod
    def _extract_from_txt(content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            return content.decode('utf-8')
        except:
            try:
                # Try other encodings
                return content.decode('latin-1')
            except:
                return str(content, errors='ignore')
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text"""
        if not text:
            return ""
        
        # Replace common encoding issues
        replacements = {
            '\xa0': ' ',  # Non-breaking space
            '\x0c': '\n',  # Form feed
            '\r\n': '\n',  # Windows line endings
            '\t': ' ',     # Tabs
            '  ': ' ',     # Double spaces
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove non-printable characters (except newlines)
        text = re.sub(r'[^\x20-\x7E\n\r]', '', text)
        
        return text.strip()
    
    @staticmethod
    def validate_file(file, filename: str) -> Tuple[bool, str]:
        """
        Validate uploaded file
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size (max 5MB)
        max_size = 5 * 1024 * 1024  # 5MB
        
        try:
            # Get file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            if file_size > max_size:
                return False, f"File too large ({file_size/1024/1024:.1f}MB). Maximum size is 5MB."
            
            if file_size == 0:
                return False, "File is empty."
            
            # Check file extension
            valid_extensions = ['.pdf', '.docx', '.doc', '.txt']
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext not in valid_extensions:
                return False, f"Invalid file type. Supported formats: {', '.join(valid_extensions)}"
            
            return True, ""
            
        except Exception as e:
            return False, f"File validation error: {str(e)}"